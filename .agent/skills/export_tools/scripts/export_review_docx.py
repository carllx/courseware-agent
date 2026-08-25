#!/usr/bin/env python3
"""
带视觉标记的人工审阅 Word 文档 (Review Document Exporter)

直接从脚本 .md 生成格式化 Word 文档，供人工阅读检查脚本与视觉素材的对应关系。
- 正文 → 黑色常规
- > [VISUAL] 块 → 红色加粗（显示 Slide ID + Layout）
- > [ACTIVITY] 块 → 蓝色加粗（显示 Type + Duration）
- 知识标签 → 灰色斜体

用法:
    python export_review_docx.py --course "实习指导" --all
    python export_review_docx.py --course "实习指导" --file S01_Mobilization
"""

import os
import sys
import argparse

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..', 'scripts', 'core')))
from script_parser import (
    parse_script, BlockType, strip_markdown,
    get_workspace_root, get_scripts_dir, list_script_files,
)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ 需要 python-docx 库。请运行: pip install python-docx")
    sys.exit(1)

# 颜色定义
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_RED = RGBColor(0xCC, 0x00, 0x00)
COLOR_BLUE = RGBColor(0x00, 0x55, 0xCC)
COLOR_GRAY = RGBColor(0x88, 0x88, 0x88)
FONT_NAME = "Microsoft YaHei"
FONT_SIZE = Pt(12)


def generate_review_docx(script_path: str, output_path: str):
    """从脚本 .md 生成审阅 Word 文档。"""
    blocks = parse_script(script_path)
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE

    basename = os.path.splitext(os.path.basename(script_path))[0]

    # 标题
    heading = doc.add_heading(basename, level=1)
    for run in heading.runs:
        run.font.name = FONT_NAME

    for b in blocks:
        if b.block_type == BlockType.HEADER:
            level = b.metadata.get("level", 2)
            # 限制层级 1-4
            level = min(max(level, 1), 4)
            h = doc.add_heading(b.content, level=level)
            for run in h.runs:
                run.font.name = FONT_NAME

        elif b.block_type == BlockType.SPEECH:
            text = strip_markdown(b.content).strip()
            if text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.color.rgb = COLOR_BLACK
                run.font.size = FONT_SIZE
                run.font.name = FONT_NAME

        elif b.block_type == BlockType.VISUAL:
            sid = b.metadata.get("slide_id", "?")
            layout = b.metadata.get("layout", "?")
            scene = b.metadata.get("scene", "")
            marker = f"[VISUAL] Slide: {sid} | Layout: {layout}"
            if scene:
                marker += f"\n  Scene: {scene}"

            p = doc.add_paragraph()
            run = p.add_run(marker)
            run.font.color.rgb = COLOR_RED
            run.font.bold = True
            run.font.size = FONT_SIZE
            run.font.name = FONT_NAME

        elif b.block_type == BlockType.SLIDE_REF:
            sid = b.metadata.get("slide_ref_id", "?")
            desc = b.content
            marker = f"[SLIDE REF] {sid}"
            if desc:
                marker += f"\n  {desc}"

            p = doc.add_paragraph()
            run = p.add_run(marker)
            run.font.color.rgb = COLOR_RED
            run.font.bold = True
            run.font.size = FONT_SIZE
            run.font.name = FONT_NAME

        elif b.block_type == BlockType.ACTIVITY:
            atype = b.metadata.get("activity_type", "?")
            duration = b.metadata.get("duration_raw", "?")
            desc = b.metadata.get("desc", "")
            marker = f"[ACTIVITY] Type: {atype} | Duration: {duration}"
            if desc:
                marker += f"\n  {desc}"

            p = doc.add_paragraph()
            run = p.add_run(marker)
            run.font.color.rgb = COLOR_BLUE
            run.font.bold = True
            run.font.size = FONT_SIZE
            run.font.name = FONT_NAME

        elif b.block_type == BlockType.TAG:
            tag_name = b.metadata.get("tag_name", "?")
            # 去掉标签名，保留内容
            content = b.content.replace(f"[{tag_name}]", "").strip()
            marker = f"[{tag_name}] {content}"

            p = doc.add_paragraph()
            run = p.add_run(marker)
            run.font.color.rgb = COLOR_GRAY
            run.font.italic = True
            run.font.size = FONT_SIZE
            run.font.name = FONT_NAME

        elif b.block_type == BlockType.SEPARATOR:
            # 添加一条水平线（通过空段落模拟）
            p = doc.add_paragraph()
            p.add_run("─" * 40).font.color.rgb = COLOR_GRAY

        elif b.block_type == BlockType.META:
            p = doc.add_paragraph()
            run = p.add_run(b.content)
            run.font.color.rgb = COLOR_GRAY
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.name = FONT_NAME

    # 保存
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"✅ 已导出: {os.path.basename(output_path)}")


def main():
    parser = argparse.ArgumentParser(description="导出带标记的审阅 Word 文档")
    parser.add_argument("--course", required=True, help="课程目录名")
    parser.add_argument("--file", help="指定单个脚本（不含扩展名）")
    parser.add_argument("--all", action="store_true", help="处理所有脚本")
    args = parser.parse_args()

    if not args.file and not args.all:
        print("❌ 请指定 --file 或 --all")
        sys.exit(1)

    workspace = get_workspace_root()
    scripts_dir = get_scripts_dir(workspace, args.course)
    output_dir = os.path.join(workspace, args.course, "delivery", "review")

    if not os.path.exists(scripts_dir):
        print(f"❌ 脚本目录不存在: {scripts_dir}")
        sys.exit(1)

    if args.all:
        files = list_script_files(scripts_dir)
    else:
        fname = args.file if args.file.endswith(".md") else f"{args.file}.md"
        files = [fname]

    for fname in files:
        fpath = os.path.join(scripts_dir, fname)
        if not os.path.exists(fpath):
            print(f"❌ 文件不存在: {fpath}")
            continue
        base = os.path.splitext(fname)[0]
        out_path = os.path.join(output_dir, f"{base}_review.docx")
        generate_review_docx(fpath, out_path)

    print(f"\n📁 输出目录: {output_dir}")


if __name__ == "__main__":
    main()
