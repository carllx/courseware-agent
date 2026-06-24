import os
import yaml
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def load_yaml(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return yaml.safe_load(f)

def create_grading_docx(course_name, course_yaml, output_path):
    doc = Document()
    
    # Setup document styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run('广州南方学院期末考查评分标准')
    run.font.size = Pt(16)
    run.font.bold = True
    
    p_sub = doc.add_paragraph('（     级                   专业2025～2026 学年度 第 2 学期）')
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'课程名称: {course_name}')
    doc.add_paragraph('考核方式: 实操类 / 演示类')
    
    sections = course_yaml['exams']['final_exam'][0]['sections']
    
    # Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '评分项'
    hdr_cells[2].text = '分值'
    
    for i, section in enumerate(sections):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = section['section_name']
        row_cells[2].text = str(section['total_score'])
        
    total_row = table.add_row().cells
    total_row[0].text = '总分'
    total_row[1].text = ''
    total_row[2].text = '100'
    
    doc.add_paragraph('\n学生课程考查大作业以“思想性、科学性、原创性”为思想导向，确保学生作业方案主题明确、内容积极健康向上；作品中所使用的主素材必须是作者原创，个别素材使用公共素材的，必须注明出处，否则视为作弊。在此基础上，以下具体评分项为本课程评分标准：\n')
    
    num_map = ['一', '二', '三', '四', '五', '六', '七', '八']
    
    ab_examples_a = []
    ab_examples_b = []
    
    for i, section in enumerate(sections):
        p = doc.add_paragraph()
        run = p.add_run(f"{num_map[i]}、{section['section_name']}（{section['total_score']}分）")
        run.font.bold = True
        
        content = section['questions'][0]['content']
        lines = [line.strip() for line in content.split('\n') if line.strip()]
        
        for line in lines:
            if '【A卷示例】' in line:
                ab_examples_a.append(line)
                doc.add_paragraph(line)
            elif '【B卷示例】' in line:
                ab_examples_b.append(line)
                doc.add_paragraph(line)
            elif line.startswith('考核要求'):
                pass
            else:
                doc.add_paragraph(line)
                
    doc.add_paragraph('\n参考示例')
    
    p = doc.add_paragraph()
    run = p.add_run('A卷参考示例                                          B卷参考示例')
    run.font.bold = True
    
    # Extract examples from practice_ab if present, or use the ones gathered
    practice_ab = course_yaml['exams']['final_exam'][0].get('practice_paper', {}).get('ab_versions', {})
    
    a_desc = practice_ab.get('A', {}).get('practice_theme', '请参考 A卷 要求设计。').strip()
    b_desc = practice_ab.get('B', {}).get('practice_theme', '请参考 B卷 要求设计。').strip()
    
    # We will just append the detailed themes or the exact requirements
    doc.add_paragraph("【A卷具体方向】\n" + a_desc)
    doc.add_paragraph("【B卷具体方向】\n" + b_desc)

    doc.save(output_path)
    print(f"Generated {output_path}")

def main():
    base_dir = '/Users/yamlam/Downloads/2025-2026-2 课程'
    
    courses = [
        ('信息可视化', '信息可视化_期末考查评分标准.docx'),
        ('交互产品开发', '交互产品开发_期末考查评分标准.docx')
    ]
    
    for c_name, out_file in courses:
        yaml_path = os.path.join(base_dir, c_name, 'course_assessment.yaml')
        out_path = os.path.join(base_dir, out_file)
        
        if os.path.exists(yaml_path):
            try:
                cyaml = load_yaml(yaml_path)
                create_grading_docx(c_name, cyaml, out_path)
            except Exception as e:
                print(f"Error processing {c_name}: {e}")
        else:
            print(f"YAML not found: {yaml_path}")

if __name__ == '__main__':
    main()
